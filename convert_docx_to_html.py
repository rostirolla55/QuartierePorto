import os
import re
import unicodedata 
import sys
from docx import Document

# --- CONFIGURAZIONI GLOBALI ---
# CORRETTO: La variabile è DOCX_DIR, non DOCS_DA_CONVERTIRE
DOCX_DIR = "DOCS_DA_CONVERTIRE"
HTML_OUTPUT_DIR = "HTML_OUTPUT"
ASSETS_BASE_DIR = "Assets/images" 

# Pattern per il marker di divisione
SPLIT_BLOCK_PATTERN = r'\[SPLIT_BLOCK:\s*(.+?\.(?:jpg|jpeg|png|gif|bmp))\]'

# --- FUNZIONI DI SUPPORTO ---
def sanitize_text(text):
    """Pulisce il testo e fa l'escape dei caratteri HTML e rimuove i \n letterali."""
    content = text
    content = unicodedata.normalize('NFC', content) 
    # Pulizia artefatti comuni
    content = content.replace('\xa0', ' ').replace('…', '...').replace('–', '-').replace('—', '-')
    content = content.replace('“', '"').replace('”', '"').replace('‘', "'").replace('’', "'") 
    
    # 🚨 AGGIUNTA IMPORTANTE: Rimuove la sequenza di caratteri letterali \n 
    # che non sono interpretati come newline ma come testo.
    content = content.replace('\\n', '') 
    
    # Escape HTML
    content = content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;') 
    content = re.sub(r'[\x00-\x1F\x7F]', '', content)
    return content

def docx_to_html(paragraph, page_id):
    """Converte il paragrafo e sostituisce il marker con il tag <img>."""
    text_buffer = []
    for run in paragraph.runs:
        current_text = sanitize_text(run.text)
        if run.bold:
            current_text = f"<strong>{current_text}</strong>"
        if run.italic:
            current_text = f"<em>{current_text}</em>"
        text_buffer.append(current_text)

    raw_text = "".join(text_buffer)
    
    # 2. Sostituzione del Marker Immagine
    def replace_marker(match):
        image_filename = match.group(1).strip()
        
        # SOLUZIONE PER IL PATH: crea path nativo e lo standardizza in URL-format
        img_path_native = os.path.join(ASSETS_BASE_DIR, page_id.lower(), image_filename)
        img_src = img_path_native.replace('\\', '/')
        
        return f'<img src="{img_src}" alt="{image_filename}">'

    html_with_images = re.sub(SPLIT_BLOCK_PATTERN, replace_marker, raw_text, flags=re.IGNORECASE)
    
    content_stripped = html_with_images.strip()
    if content_stripped:
        return f"<p>{content_stripped}</p>"
    return ""

# --- FUNZIONE PRINCIPALE DI CONVERSIONE E SPLIT ---

def convert_docx_and_split(page_id, docx_filename):
    """
    Funzione principale che carica il DOCX, lo converte e lo divide in blocchi HTML.
    """
    normalized_page_id = page_id.lower()
    
    # *** CORREZIONE: Usa DOCX_DIR che è definito in alto, non DOCS_DA_CONVERTIRE ***
    docx_path = os.path.join(DOCX_DIR, docx_filename) 
    
    if not os.path.exists(docx_path):
        print(f"ERRORE: File DOCX non trovato al percorso: {docx_path}")
        return False

    os.makedirs(HTML_OUTPUT_DIR, exist_ok=True)
    
    try:
        document = Document(docx_path)
    except Exception as e:
        print(f"ERRORE: Impossibile aprire o leggere il file DOCX: {e}")
        return False

    html_blocks = [[]] 
    current_block_index = 0
    
    # Iterazione robusta sui Paragrafi (document.paragraphs)
    for paragraph in document.paragraphs:
        
        marker_match = re.search(SPLIT_BLOCK_PATTERN, paragraph.text, flags=re.IGNORECASE)
        
        if marker_match or paragraph.text.strip():
            html_content = docx_to_html(paragraph, normalized_page_id)
            
            if marker_match:
                if html_content:
                    html_blocks[current_block_index].append(html_content)
                current_block_index += 1
                html_blocks.append([]) 
            elif html_content:
                html_blocks[current_block_index].append(html_content)
        
    # --- Salvataggio dei File ---
    
    base_filename = os.path.splitext(docx_filename)[0]
    html_blocks = [block for block in html_blocks if block]
    num_blocks = len(html_blocks)
    success = True
    
    if num_blocks == 0:
        return False

    for i, block in enumerate(html_blocks):
        block_number = i + 1 
        html_filename = f"{base_filename}{block_number}.html"
        html_path = os.path.join(HTML_OUTPUT_DIR, html_filename)
        
        try:
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write("\n".join(block))
        except Exception as e:
            print(f"ERRORE GRAVE: Impossibile scrivere il file {html_path}: {e}")
            success = False

    return success

# Entry point per l'esecuzione dello script.
if __name__ == "__main__":
    if len(sys.argv) < 3:
        sys.exit(1)
        
    page_id = sys.argv[1]
    docx_filename = sys.argv[2]
    
    if convert_docx_and_split(page_id, docx_filename):
        sys.exit(0)
    else:
        sys.exit(1)